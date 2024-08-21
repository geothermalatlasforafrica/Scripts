import os
from io import BytesIO

from PIL import Image
import docx
import pandas as pd
import requests
from docx import Document
from docx.shared import Pt


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"


def add_paragraph(doc, text, line_spacing=1.0):
    p = doc.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.line_spacing = line_spacing

    return p


def run():
    # Step 2: Read the Excel file
    df = pd.read_excel('gaa_metadata_restructured.xlsx')

    # Step 3: Extract the required columns
    df = df[['parent_group', 'filename', 'full_name', 'description', 'keywords', 'source', 'date', 'coverage', 'resolution']]

    # Order the dataframe alphabetically by full_name
    df = df.sort_values(by='full_name')

    # Step 4: Create a new Word document
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    bbox = "-45,-25,40,60"

    # Step 5: Iterate over the rows of the DataFrame
    for index, row in df.iterrows():
        # Stop at row 5
        # if index == 5:
        #     break

        # Only include rows whose parent_group is "Geoscientific"
        if row['parent_group'] != 'Geoscientific':
            continue

        print(row['full_name'])
        # Add a new page
        if index != 0:
            doc.add_page_break()

        doc.add_heading(row['full_name'])

        add_paragraph(doc, 'description: ' + str(row['description']))
        add_paragraph(doc, 'keywords: ' + str(row['keywords']))

        # If the source is a URL, add a hyperlink to it. Otherwise, use the text 'no online source available'
        if pd.isnull(row['source']):
            add_paragraph(doc, 'source: no online source available')
        else:
            p = add_paragraph(doc, 'source: ')
            add_hyperlink(p, 'link', row['source'])

        add_paragraph(doc, 'date: ' + str(row['date']))
        add_paragraph(doc, 'coverage: ' + str(row['coverage']))

        # Add a GeoServer WMS request for a JPG image of Africa
        file_name_without_ext = os.path.splitext(row['filename'])[0]
        wms_url = f"https://gaa-proxy.azurewebsites.net/geoserver/gaa-dev/wms?SERVICE=WMS&VERSION=1.3.0&REQUEST=GetMap&FORMAT=image%2Fpng&TRANSPARENT=true&STYLES&LAYERS=gaa-dev%3A{file_name_without_ext}&exceptions=application%2Fvnd.ogc.se_inimage&SRS=EPSG%3A4326&WIDTH=747&HEIGHT=768&BBOX={bbox}"
        response = requests.get(wms_url)

        # Check if the request was successful
        if response.status_code != 200:
            raise Exception('Failed to fetch image from GeoServer')

        # Add a OpenStreetMap WMS request for a PNG image of Africa
        osm_url = f"https://ows.terrestris.de/osm/service?SERVICE=WMS&VERSION=1.3.0&REQUEST=GetMap&FORMAT=image%2Fpng&TRANSPARENT=true&LAYERS=OSM-WMS&STYLES&TILED=false&WIDTH=747&HEIGHT=768&CRS=EPSG%3A4326&BBOX={bbox}"
        response_osm = requests.get(osm_url)

        # Check if the request was successful
        if response_osm.status_code != 200:
            raise Exception('Failed to fetch image from OpenStreetMap')

        img1 = Image.open(BytesIO(response.content)).convert("RGBA")
        img2 = Image.open(BytesIO(response_osm.content)).convert("RGBA")

        # Ensure the images are the same size
        if img1.size != img2.size:
            img2 = img2.resize(img1.size)

        # Blend the images
        alpha = 0.5  # set your own alpha
        img_blend = Image.blend(img1, img2, alpha)

        # Save the result to a BytesIO stream
        image_stream_blend = BytesIO()
        img_blend.save(image_stream_blend, format='PNG')

        # Get the legend graphic
        legend_url = f"https://gaa-proxy.azurewebsites.net/geoserver/wms?REQUEST=GetLegendGraphic&VERSION=1.0.0&FORMAT=image/png&WIDTH=20&HEIGHT=20&STRICT=false&style=gaa-dev:{file_name_without_ext}_style"
        response_legend = requests.get(legend_url)

        # Check if the request was successful
        if response_legend.status_code != 200:
            raise Exception('Failed to fetch legend from GeoServer')

        # Open the third image
        img3 = Image.open(BytesIO(response_legend.content)).convert("RGBA")

        # Resize the third image to the desired size for the overlay
        # overlay_size = (200, 200)  # set your own overlay size
        # img3 = img3.resize(overlay_size)

        # Paste the third image onto the blended image in the bottom left corner
        bottom_left_corner = (0, img_blend.height - img3.height)
        img_blend.paste(img3, bottom_left_corner, img3)

        # Save the result to a BytesIO stream
        image_stream_blend = BytesIO()
        img_blend.save(image_stream_blend, format='PNG')

        # Add the resulting image to the Word document
        image_stream_blend.seek(0)  # Go to the start of the BytesIO stream
        doc.add_picture(image_stream_blend, width=docx.shared.Inches(5.3))

    # Step 6: Save the Word document
    doc.save('output.docx')


if __name__ == '__main__':
    run()
